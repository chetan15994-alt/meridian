"""ATS-safe .docx rendering of a tailored resume. Single-column, standard headings,
no tables/graphics — maximally parseable by applicant tracking systems. Faithful to a
conventional professional CV structure with subtle typographic polish."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INK   = RGBColor(0x1a, 0x1a, 0x1a)
NAVY  = RGBColor(0x1f, 0x3a, 0x5f)
GREY  = RGBColor(0x55, 0x55, 0x55)

def _heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    pf = p.paragraph_format
    pf.space_before = Pt(8); pf.space_after = Pt(2)
    # thin rule under the heading via bottom border
    _bottom_border(p)
    return p

def _bottom_border(paragraph):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1');    bottom.set(qn('w:color'), '1f3a5f')
    pbdr.append(bottom); pPr.append(pbdr)

def _fmt_education(ed):
    """Render an education entry whether it's a structured dict or a plain string."""
    if isinstance(ed, dict):
        deg   = (ed.get("degree") or "").strip()
        field = (ed.get("field") or "").strip()
        inst  = (ed.get("institution") or ed.get("school") or "").strip()
        year  = (ed.get("year") or "").strip()
        left  = deg + (f", {field}" if field else "")
        parts = [p for p in [left, inst] if p]
        s = "  —  ".join(parts) if parts else ""
        if year: s += f"  ({year})"
        return s.strip()
    return str(ed).strip()

def _fmt_cert(c):
    if isinstance(c, dict):
        name = (c.get("name") or "").strip()
        issuer = (c.get("issuer") or "").strip()
        year = (c.get("year") or "").strip()
        s = name + (f" — {issuer}" if issuer else "")
        if year: s += f" ({year})"
        return s.strip()
    return str(c).strip()

def render_docx(tailored, outdir):
    os.makedirs(outdir, exist_ok=True)
    doc = Document()

    # Tighter, professional margins
    for section in doc.sections:
        section.top_margin = Pt(40); section.bottom_margin = Pt(40)
        section.left_margin = Pt(54); section.right_margin = Pt(54)

    base = doc.styles["Normal"].font
    base.name = "Calibri"; base.size = Pt(10.5); base.color.rgb = INK

    # ── Header: name + title + contact ──
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run(tailored.get("name", "")); r.bold = True
    r.font.size = Pt(20); r.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(0)

    if tailored.get("title"):
        t = doc.add_paragraph(); tr = t.add_run(tailored["title"])
        tr.font.size = Pt(11); tr.bold = True; tr.font.color.rgb = GREY
        t.paragraph_format.space_after = Pt(2)

    c = tailored.get("contact", {}) or {}
    contact_bits = [c.get("email"), c.get("phone"), c.get("linkedin"), c.get("location")]
    contact = "  |  ".join(b for b in contact_bits if b and not str(b).startswith("[fill"))
    if contact:
        cp = doc.add_paragraph(); cr = cp.add_run(contact)
        cr.font.size = Pt(9); cr.font.color.rgb = GREY
        cp.paragraph_format.space_after = Pt(2)

    # ── Professional Summary ──
    if tailored.get("summary"):
        _heading(doc, "Professional Summary")
        p = doc.add_paragraph(tailored["summary"])
        p.paragraph_format.space_after = Pt(2)

    # ── Core Competencies ──
    if tailored.get("skills"):
        _heading(doc, "Core Competencies")
        p = doc.add_paragraph(" · ".join(tailored["skills"]))
        p.paragraph_format.space_after = Pt(2)

    # ── Professional Experience ──
    if tailored.get("experiences"):
        _heading(doc, "Professional Experience")
        for e in tailored["experiences"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(f"{e.get('company','')} — {e.get('role','')}")
            rr.bold = True; rr.font.size = Pt(11); rr.font.color.rgb = INK
            if e.get("dates"):
                rr2 = p.add_run(f"    ({e['dates']})")
                rr2.italic = True; rr2.font.size = Pt(9); rr2.font.color.rgb = GREY
            for b in e.get("bullets", []):
                bp = doc.add_paragraph(b, style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)

    # ── Education ──
    edus = tailored.get("education") or []
    if edus:
        _heading(doc, "Education")
        for ed in edus:
            line = _fmt_education(ed)
            if line:
                ep = doc.add_paragraph(line); ep.paragraph_format.space_after = Pt(0)

    # ── Certifications ──
    certs = tailored.get("certifications") or []
    if certs:
        _heading(doc, "Certifications")
        for cert in certs:
            line = _fmt_cert(cert)
            if line:
                doc.add_paragraph(line, style="List Bullet")

    import datetime as _dt
    safe = "".join(ch for ch in tailored.get("tailored_for", "resume")
                   if ch.isalnum() or ch in " _-")[:50].strip().replace(" ", "_")
    name_slug = "".join(ch for ch in tailored.get("name", "Resume") if ch.isalnum() or ch == " ").strip().replace(" ", "_")
    stamp = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(outdir, f"{name_slug}_{safe}_{stamp}.docx")
    doc.save(path)
    return path

def parse_back_check(path):
    """Re-open the .docx and extract text to confirm it parses cleanly (ATS sanity check)."""
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    n = len(text.strip())
    ok = n > 100 and "{" not in text and "}" not in text   # no leaked dict/JSON
    return ok, n

def render_cover_letter(tailored, outdir):
    cl = tailored.get("cover_letter", "")
    if not cl: return None
    os.makedirs(outdir, exist_ok=True)
    import datetime as _dt
    safe = "".join(ch for ch in tailored.get("tailored_for", "cover")
                   if ch.isalnum() or ch in " _-")[:50].strip().replace(" ", "_")
    stamp = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(outdir, f"CoverLetter_{safe}_{stamp}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(cl)
    return path
