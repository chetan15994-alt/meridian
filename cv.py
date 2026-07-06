"""CV upload parsing, versioning and completeness.

Three responsibilities:
  1. Parse  — extract structured YAML from a PDF or DOCX upload via Claude API
  2. Version — snapshot resume_master.yaml before every save; restore on demand
  3. Check  — scan for placeholders, empty required fields, weak bullets

All filesystem ops are self-contained (no Streamlit); pure-testable."""
import os, re, json, yaml, shutil, datetime, base64, io

HERE       = os.path.dirname(__file__)
MASTER     = os.path.join(HERE, "resume_master.yaml")
VERSIONS   = os.path.join(HERE, "cv_versions")
INDEX_FILE = os.path.join(VERSIONS, "index.json")

os.makedirs(VERSIONS, exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — SCHEMA + PARSE PROMPT
# ─────────────────────────────────────────────────────────────────────────────

# Meridian's resume_master.yaml schema (canonical)
SCHEMA = """
{
  "name": "Full Name",
  "title": "Professional headline e.g. Senior Product Manager",
  "contact": {
    "email": "email@example.com",
    "phone": "+91 ...",
    "linkedin": "https://linkedin.com/in/...",
    "location": "City, Country"
  },
  "summary": "2-4 sentence professional summary in first person",
  "skills": ["skill1", "skill2"],
  "experiences": [
    {
      "company": "Company Name",
      "role": "Job Title",
      "dates": "Mon YYYY – Mon YYYY  (use Present for current)",
      "bullets": [
        "Strong action-verb led bullet with a measurable outcome."
      ]
    }
  ],
  "education": [
    {
      "institution": "Institution Name",
      "degree": "Degree e.g. MBA",
      "field": "Field e.g. Marketing",
      "year": "YYYY"
    }
  ],
  "certifications": ["Certification name and issuer"]
}
"""

PARSE_PROMPT = (
    "You are an expert CV parser. Extract every piece of information from the CV/resume below "
    "into the EXACT JSON schema provided. Rules:\n"
    "- Extract ONLY what is written. Never invent, assume, or add information.\n"
    "- Preserve the original wording of bullet points (light formatting cleanup is fine).\n"
    "- Dates: 'Mon YYYY' format. Use 'Present' for the current role's end date.\n"
    "- Skills: individual items, not grouped categories.\n"
    "- Bullets: one string per bullet, no leading dash/hyphen.\n"
    "- If a field is absent from the CV, use null (not a placeholder string).\n"
    "- Return ONLY the JSON object — no markdown fences, no explanation.\n\n"
    "SCHEMA:\n" + SCHEMA
)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — EXTRACT TEXT
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_docx(file_bytes: bytes) -> str:
    """Return plain text from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t: lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text: lines.append(row_text)
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — PARSE via Claude API
# ─────────────────────────────────────────────────────────────────────────────

def _call_anthropic_parse(content_blocks: list, base_url: str, model: str,
                           api_key: str, max_tokens: int = 3000) -> dict:
    """Send content_blocks to Anthropic /v1/messages and return parsed JSON dict."""
    import requests as req
    url = base_url.rstrip("/")
    url = url + "/messages" if url.endswith("/v1") else url + "/v1/messages"
    headers = {"x-api-key": api_key, "anthropic-version": "2023-06-01",
                "content-type": "application/json"}
    payload = {"model": model, "max_tokens": max_tokens,
                "messages": [{"role": "user", "content": content_blocks}]}
    r = req.post(url, json=payload, headers=headers, timeout=120)
    if r.status_code == 401:
        raise ValueError("Anthropic 401 — check your API key in Settings.")
    r.raise_for_status()
    text = "".join(b.get("text","") for b in r.json().get("content",[]) if b.get("type")=="text")
    # strip possible markdown fences
    clean = re.sub(r"```(?:json)?\s*", "", text).replace("```","").strip()
    a, b = clean.find("{"), clean.rfind("}")
    if a == -1 or b == -1:
        raise ValueError(f"Model did not return valid JSON. Got: {text[:200]}")
    return json.loads(clean[a:b+1])


def parse_from_docx(file_bytes: bytes, base_url: str, model: str, api_key: str) -> dict:
    """Extract text from DOCX then ask Claude to parse it."""
    text = extract_text_docx(file_bytes)
    if not text.strip():
        raise ValueError("Could not extract text from DOCX — file may be image-only or corrupt.")
    content = [{"type": "text", "text": PARSE_PROMPT + "\n\nCV TEXT:\n" + text}]
    return _call_anthropic_parse(content, base_url, model, api_key)


def parse_from_pdf(file_bytes: bytes, base_url: str, model: str, api_key: str) -> dict:
    """Send PDF as a base64 document block — Claude reads it natively."""
    b64 = base64.standard_b64encode(file_bytes).decode()
    content = [
        {"type": "document",
         "source": {"type": "base64", "media_type": "application/pdf", "data": b64}},
        {"type": "text", "text": PARSE_PROMPT},
    ]
    return _call_anthropic_parse(content, base_url, model, api_key)


def parsed_to_yaml(parsed: dict) -> dict:
    """Normalise the Claude-parsed dict to match Meridian's YAML schema exactly."""
    def _str(v, default=""): return str(v).strip() if v else default
    def _list(v): return [str(i).strip() for i in v if i] if isinstance(v, list) else []

    contact_raw = parsed.get("contact") or {}
    exps_raw    = parsed.get("experiences") or parsed.get("experience") or []
    edu_raw     = parsed.get("education") or []

    experiences = []
    for e in exps_raw:
        if not isinstance(e, dict): continue
        bullets = _list(e.get("bullets") or [])
        experiences.append({
            "company": _str(e.get("company") or e.get("employer")),
            "role":    _str(e.get("role") or e.get("title")),
            "dates":   _str(e.get("dates") or
                           (f"{e.get('start','')} – {e.get('end','')}").strip(" –")),
            "bullets": bullets or ["[fill in: add bullet points for this role]"],
        })

    education = []
    for ed in edu_raw:
        if not isinstance(ed, dict): continue
        education.append({
            "institution": _str(ed.get("institution") or ed.get("school")),
            "degree":      _str(ed.get("degree")),
            "field":       _str(ed.get("field") or ed.get("major")),
            "year":        _str(ed.get("year") or ed.get("graduation_year")),
        })

    return {
        "name":           _str(parsed.get("name")),
        "title":          _str(parsed.get("title") or parsed.get("headline")),
        "contact": {
            "email":    _str(contact_raw.get("email") or parsed.get("email")),
            "phone":    _str(contact_raw.get("phone") or parsed.get("phone")),
            "linkedin": _str(contact_raw.get("linkedin") or parsed.get("linkedin")),
            "location": _str(contact_raw.get("location") or parsed.get("location")),
        },
        "summary":        _str(parsed.get("summary")),
        "skills":         _list(parsed.get("skills")),
        "experiences":    experiences,
        "education":      education,
        "certifications": _list(parsed.get("certifications")),
    }


def validate_resume(data):
    """Return (ok, problems[]). A resume must at least be a dict with name/experiences
    so the renderer never crashes on a malformed parse or corrupt snapshot."""
    problems = []
    if not isinstance(data, dict):
        return False, ["not a valid resume object"]
    if not str(data.get("name") or "").strip():
        problems.append("missing name")
    exps = data.get("experiences")
    if not isinstance(exps, list):
        problems.append("experiences is not a list")
    else:
        for i, e in enumerate(exps):
            if not isinstance(e, dict) or not e.get("company"):
                problems.append(f"experience[{i}] malformed")
    # education/certs must be lists if present
    for f in ("education", "certifications", "skills"):
        if f in data and not isinstance(data[f], list):
            problems.append(f"{f} should be a list")
    ok = not any(p for p in problems if "malformed" in p or "not a" in p or "missing name" in p)
    return ok, problems


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — VERSIONING
# ─────────────────────────────────────────────────────────────────────────────

def _load_index() -> list:
    import fileio
    try:
        return fileio.read_json(INDEX_FILE, default=[]) or []
    except Exception:
        return []

def _save_index(idx: list):
    import fileio
    fileio.write_json_atomic(INDEX_FILE, idx)

def save_snapshot(label: str = "", source: str = "manual") -> str:
    """Snapshot the current resume_master.yaml before any destructive save.
    Returns the snapshot filename, or '' if master doesn't exist yet."""
    if not os.path.exists(MASTER): return ""
    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:20]  # yyyymmdd_HHMMSS_μμμ
    fn   = f"resume_{ts}.yaml"
    dest = os.path.join(VERSIONS, fn)
    shutil.copy2(MASTER, dest)
    idx  = _load_index()
    idx.insert(0, {
        "filename":  fn,
        "timestamp": datetime.datetime.now().isoformat(timespec="seconds"),
        "label":     label or "",
        "source":    source,
    })
    _save_index(idx)
    return fn

def list_snapshots() -> list:
    """Return list of snapshot metadata dicts, newest first."""
    idx = _load_index()
    out = []
    for entry in idx:
        fn   = entry.get("filename","")
        path = os.path.join(VERSIONS, fn)
        entry["exists"] = os.path.exists(path)
        out.append(entry)
    return out

def restore_snapshot(filename: str):
    """Overwrite resume_master.yaml with the given snapshot (auto-snapshots current first)."""
    src = os.path.join(VERSIONS, filename)
    if not os.path.exists(src):
        raise FileNotFoundError(f"Snapshot not found: {filename}")
    # Read target content BEFORE auto-snapshot (avoid overwriting src with itself)
    target_content = open(src, "rb").read()
    # validate the snapshot before clobbering the live master
    try:
        parsed = yaml.safe_load(target_content)
        ok, problems = validate_resume(parsed)
        if not ok:
            raise ValueError("snapshot is corrupt: " + "; ".join(problems))
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"snapshot is unreadable: {e}")
    save_snapshot(label=f"before-restore-{filename[:15]}", source="auto")
    with open(MASTER, "wb") as f:
        f.write(target_content)
    # the resume just changed on disk — clear every cache derived from it
    import settings
    settings.invalidate_resume_caches()

def delete_snapshot(filename: str):
    """Delete a snapshot file and remove it from the index."""
    path = os.path.join(VERSIONS, filename)
    if os.path.exists(path): os.remove(path)
    idx = [e for e in _load_index() if e.get("filename") != filename]
    _save_index(idx)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — COMPLETENESS + QUALITY CHECK
# ─────────────────────────────────────────────────────────────────────────────

_PLACEHOLDER_RE = re.compile(r"\[fill in[^\]]*\]", re.I)
_WEAK_VERBS     = {"responsible for","worked on","helped","assisted","involved in",
                   "participated in","contributed to","supported","various","etc"}
_ACTION_VERBS   = {"led","built","designed","launched","owned","drove","scaled","reduced",
                   "increased","created","delivered","shipped","grew","managed","defined",
                   "architected","improved","spearheaded","implemented","established",
                   "pioneered","transformed","secured","negotiated","generated","optimized"}

def check_completeness(data: dict) -> list:
    """Scan resume dict for issues. Returns list of {level, field, message} dicts.
    level: 'error' (will break output) | 'warning' (degrades quality) | 'tip'"""
    issues = []

    def _err(field, msg): issues.append({"level":"error",  "field":field, "message":msg})
    def _wrn(field, msg): issues.append({"level":"warning","field":field, "message":msg})
    def _tip(field, msg): issues.append({"level":"tip",    "field":field, "message":msg})

    raw_yaml = yaml.safe_dump(data, allow_unicode=True)

    # Placeholders — errors (ship directly into output)
    for ph in _PLACEHOLDER_RE.findall(raw_yaml):
        _err("placeholder", f"Placeholder found: '{ph}' — fill this in before tailoring")

    # Required top-level fields
    for field in ("name","title","summary"):
        val = str(data.get(field) or "").strip()
        if not val:
            _err(field, f"'{field}' is empty — required for every tailored resume")
        elif len(val) < 15:
            _wrn(field, f"'{field}' looks very short ({len(val)} chars) — add more detail")

    # Contact
    contact = data.get("contact") or {}
    for cf in ("email","linkedin","location"):
        if not str(contact.get(cf) or "").strip():
            _wrn(f"contact.{cf}", f"Contact '{cf}' is missing")

    # Skills
    skills = data.get("skills") or []
    if len(skills) < 5:
        _wrn("skills", f"Only {len(skills)} skills listed — aim for 10–15 for better keyword matching")

    # Experience
    exps = data.get("experiences") or []
    if not exps:
        _err("experiences", "No experience entries — this is required")
    for i, ex in enumerate(exps):
        label = ex.get("company","?") + " / " + ex.get("role","?")
        bullets = ex.get("bullets") or []
        if not ex.get("dates","").strip():
            _wrn(f"exp[{i}].dates", f"{label}: dates are missing")
        if len(bullets) == 0:
            _err(f"exp[{i}].bullets", f"{label}: no bullets — required")
        elif len(bullets) < 2:
            _wrn(f"exp[{i}].bullets", f"{label}: only 1 bullet — aim for 3-5")
        for j, b in enumerate(bullets):
            bl = str(b or "").lower().strip()
            # Weak openers
            if any(bl.startswith(w) for w in _WEAK_VERBS):
                _wrn(f"exp[{i}].bullets[{j}]",
                     f"{label}: bullet starts with weak phrase — rewrite with action verb")
            # No metrics
            if not re.search(r"\d", b or ""):
                _tip(f"exp[{i}].bullets[{j}]",
                     f"{label}: bullet has no number/metric — quantify the outcome if possible")

    # Education
    edu = data.get("education") or []
    if not edu:
        _wrn("education", "No education entries")

    return issues


def completeness_summary(issues: list) -> dict:
    """Roll up issue counts by level."""
    return {
        "errors":   sum(1 for i in issues if i["level"]=="error"),
        "warnings": sum(1 for i in issues if i["level"]=="warning"),
        "tips":     sum(1 for i in issues if i["level"]=="tip"),
        "score":    max(0, 100 - sum(1 for i in issues if i["level"]=="error")*15
                               - sum(1 for i in issues if i["level"]=="warning")*5
                               - sum(1 for i in issues if i["level"]=="tip")),
    }


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — DIFF (master vs snapshot)
# ─────────────────────────────────────────────────────────────────────────────

def diff_snapshots(fn_a: str, fn_b: str = "current") -> list:
    """Simple field-by-field diff of two snapshots (or current vs a snapshot).
    Returns list of {field, old, new} for changed scalar fields."""
    def _load(fn):
        import fileio
        if fn == "current":
            return fileio.read_yaml(MASTER, default={}) or {}
        path = os.path.join(VERSIONS, fn)
        return fileio.read_yaml(path, default={}) or {}

    a, b = _load(fn_a), _load(fn_b)
    changes = []

    def _compare(va, vb, path=""):
        if isinstance(va, dict) and isinstance(vb, dict):
            for k in set(list(va.keys()) + list(vb.keys())):
                _compare(va.get(k), vb.get(k), f"{path}.{k}" if path else k)
        elif isinstance(va, list) and isinstance(vb, list):
            if va != vb:
                changes.append({"field": path,
                                 "old": f"{len(va)} items",
                                 "new": f"{len(vb)} items"})
        else:
            sv, sv2 = str(va or ""), str(vb or "")
            if sv != sv2:
                changes.append({"field": path,
                                 "old": sv[:120] + ("…" if len(sv) > 120 else ""),
                                 "new": sv2[:120] + ("…" if len(sv2) > 120 else "")})
    _compare(a, b)
    return changes
